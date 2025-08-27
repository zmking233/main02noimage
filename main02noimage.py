
import sys
import os
import subprocess
import traceback
import re
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QListWidget, QLabel, QFileDialog, QPushButton, QSpacerItem, QSizePolicy
)
from PyQt5.QtCore import Qt
from docx import Document
import fitz  # PyMuPDF
import shutil
import base64
from bs4 import BeautifulSoup

if getattr(sys, 'frozen', False):
    base_path = os.path.dirname(sys.executable)
else:
    base_path = os.path.dirname(os.path.abspath(__file__))

# ---------------- 文书处理函数 ----------------


def special_segment_trial_result(text: str) -> str:
    # 先给所有“一、二、三、四、五、六、七、八、九、十、十一、十二、……”的序号前加换行符，避免粘连
    # 这里用负向前瞻防止重复加换行符
    text = re.sub(r'(?<!\n)(?=[一二三四五六七八九十]+、)', r'\n', text)

    # 按行拆分
    lines = text.split('\n')
    new_lines = []

    for line in lines:
        # 统计每行汉字数量
        hanzi_count = len(re.findall(r'[\u4e00-\u9fa5]', line))
        if hanzi_count <= 35:
            new_lines.append(line + "<PARA>")
        else:
            new_lines.append(line)

    return "\n".join(new_lines)

def read_docx_full_text(docx_path):
    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    full_text = "\n".join(p.text.strip() for p in paragraphs if p.text.strip())
    return full_text

def get_case_name_from_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = doc.paragraphs[:2]
    clean_paras = []
    for p in paragraphs:
        text = p.text.strip()
        if text:
            text_no_space = re.sub(r'\s+', '', text)
            cleaned = re.sub(r'[^\w\u4e00-\u9fa5，。！？、：；（）《》“”‘’—\-\.]', '', text_no_space)
            clean_paras.append(cleaned)
    case_name = ''.join(clean_paras)
    return case_name if case_name else "未知案件名称"


def get_case_number_from_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    if len(paragraphs) >= 3:
        case_number = paragraphs[2].text.strip()
        return case_number if case_number else "未知案号"
    else:
        return "未知案号"

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def parse_fields(text, case_name, case_number):
    m_process = re.search(r'([^\n]*?(审理终结|审查终结|审理了本案)[^\n]*)', text)
    litigation_process = m_process.group(1).strip() if m_process else "诉讼记录缺失"

    if m_process:
        party_text = text.split(case_number)[-1].split(m_process.group(1))[0]
    else:
        party_text = ""
    party_lines = [line.strip() for line in party_text.split('\n') if any(k in line for k in ['原告', '被告', '上诉人', '被上诉人', '再审申请人', '被申请人']) and line.strip()]
    parties_info = '\n'.join(party_lines) if party_lines else "当事人信息缺失"

    m_case_info = re.search(r'(?:审理终结。|审查终结。)(.*?)(?=本院认为|本院经审查认为|本院再审认为)', text, re.DOTALL)
    case_info = m_case_info.group(1).strip() if m_case_info else "案件基本情况缺失"

    m_analysis = re.search(r'(本院认为|本院经审查认为|本院再审认为)(.*?(判决如下：|裁定如下：))', text, re.DOTALL)
    trial_analysis = (m_analysis.group(1) + m_analysis.group(2)).strip() if m_analysis else "裁判分析过程缺失"

    m_result = re.search(r'(判决如下：|裁定如下：)(.*?)(?=审判长|审 判 长)', text, re.DOTALL)
    trial_result = m_result.group(2).strip() if m_result else "裁判结果缺失"

    m_judge_info = re.search(r'(审判长.*|审 判 长.*)', text, re.DOTALL)
    judge_info = m_judge_info.group(1).strip() if m_judge_info else "人员信息缺失"

    # 返回所有提取字段
    return {
        'case_name': case_name,
        'case_number': case_number,
        'litigation_process': litigation_process,
        'parties_info': parties_info,
        'case_info': case_info,
        'trial_analysis': trial_analysis,
        'trial_result': trial_result,
        'judge_info': judge_info
    }

# --------- PDF 特殊提取逻辑 ---------


def extract_text_from_pdf(pdf_path, output_dir, base_name):
    txt_path = os.path.join(output_dir, f"{base_name}_debug.txt")

    doc = fitz.open(pdf_path)
    all_text = ""
    for page in doc:
        all_text += page.get_text()
    doc.close()

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(all_text)

    with open(txt_path, "r", encoding="utf-8") as f:
        processed_text = f.read()

    # 先按行分割（保留换行）
    lines = [line.strip() for line in processed_text.splitlines() if line.strip()]

    # 去除行内多余空白
    lines_no_spaces = [re.sub(r'\s+', '', line) for line in lines]

    page_num_pattern = re.compile(r'^\d+/\d+$')
    filtered_lines = [line for line in lines_no_spaces if not page_num_pattern.match(line)]

    # 给符合条件的行后添加<PARA>
    lines_with_para = add_para_tags(filtered_lines)

    # 合并成字符串
    full_text_with_para = ''.join(lines_with_para)

    # 按<PARA>分割成段
    litigation_paragraphs = [p.strip() for p in full_text_with_para.split('<PARA>') if p.strip()]

    # 用换行连接成带换行的多段文本
    full_text_for_litigation = "\n".join(litigation_paragraphs)

    return extract_text_from_txt(lines_no_spaces, full_text_for_litigation)


def special_segment_trial_result(text: str) -> str:
    # 给所有“一、二、三、...”前加换行符，防止粘连
    text = re.sub(r'(?<!\n)(?=[一二三四五六七八九十]+、)', r'\n', text)
    # 按行处理，长度不超过35的行末尾加<PARA>
    lines = text.split('\n')
    new_lines = []
    for line in lines:
        hanzi_count = len(re.findall(r'[\u4e00-\u9fa5]', line))
        if hanzi_count <= 35:
            new_lines.append(line + "<PARA>")
        else:
            new_lines.append(line)
    return "\n".join(new_lines)


def extract_judge_info_from_no_spaces(txt_no_spaces):
    lines = txt_no_spaces.split('\n')

    # 找到第一个只含“审判人员”的行
    judge_start_line = -1
    for i, line in enumerate(lines):
        if line.strip() == "审判人员":
            judge_start_line = i
            break
    if judge_start_line == -1:
        return ""

    # 从judge_start_line+1开始找包含“审判长”相关的行
    start_judge = -1
    for i in range(judge_start_line + 1, len(lines)):
        if any(x in lines[i] for x in ["审判长", "审 判 长", "审判 长"]):
            start_judge = i
            break
    if start_judge == -1:
        return ""

    # 从start_judge开始找第一个包含“书记员”相关的行
    end_judge = -1
    for i in range(start_judge, len(lines)):
        if any(x in lines[i] for x in ["书记员", "书记 员", "书 记 员"]):
            end_judge = i
            break
    if end_judge == -1:
        return ""

    judge_lines = lines[start_judge:end_judge + 1]
    judge_info = "\n".join(judge_lines).strip()
    return judge_info

def add_para_tags(lines):
    """对每行判断，少于等于35个字符且以句号结尾的行，或少于10个字符的行，后加<PARA>"""
    new_lines = []
    for line in lines:
        if (len(line) <= 40 and line.endswith('。')) or len(line) < 10:
            new_lines.append(line + '<PARA>')
        else:
            new_lines.append(line)
    return new_lines

def extract_text_from_txt(txt_no_spaces, full_text_for_litigation):
    # 提取案名和案号
    first_colon_idx = next((i for i, line in enumerate(txt_no_spaces) if '：' in line), None)
    if first_colon_idx is None:
        fuzzy_title = ""
        first_colon_line = ""
    else:
        text_before_first_colon = "".join(txt_no_spaces[:first_colon_idx]).strip()

    keywords = ["判决书", "裁定书", "案"]
    indices = []
    for kw in keywords:
        idx = text_before_first_colon.find(kw)
        indices.append(idx if idx != -1 else float('inf'))

    min_index = min(indices)
    if min_index == float('inf'):
        fuzzy_title = text_before_first_colon
    else:
        # 找到最早出现的关键词
        kw_index = indices.index(min_index)
        kw = keywords[kw_index]
        # 截取包含关键词本身
        fuzzy_title = text_before_first_colon[:min_index + len(kw)].strip()

    first_colon_line = txt_no_spaces[first_colon_idx]

    before_colon, after_colon = "", ""
    if first_colon_line:
        parts = first_colon_line.split('：')
        if len(parts) >= 2:
            before_colon = parts[0]
            after_colon = parts[1]
    m_admin = re.search(r'(行政.*)', fuzzy_title)
    part_before = m_admin.group(1) if m_admin else ""
    case_name = (after_colon + part_before).strip()

    colon_count = 0
    second_colon_idx = None
    for i, line in enumerate(txt_no_spaces):
        if '：' in line:
            colon_count += 1
            if colon_count == 2:
                second_colon_idx = i
                break
    if second_colon_idx is not None:
        case_number = txt_no_spaces[second_colon_idx].split('：', 1)[1].strip()
    else:
        case_number = "未知案号"

    # 定位当事人与审理经过之间的行，提取当事人信息
    try:
        start_party = next(i for i, l in enumerate(txt_no_spaces) if '当事人' in l)
        end_party = next(i for i, l in enumerate(txt_no_spaces) if '审理经过' in l)
    except StopIteration:
        start_party, end_party = 0, 0

    parties_txt_no_spaces = []
    for line in txt_no_spaces[start_party + 1:end_party]:
        if any(k in line for k in ['原告', '被告', '上诉人', '被上诉人', '再审申请人', '被申请人']):
            if '住所地' in line:
                line = line.split('住所地')[0].strip()
            parties_txt_no_spaces.append(line.strip())
    parties_info = "\n".join(parties_txt_no_spaces).replace(fuzzy_title, "").strip()

    # 提取诉讼过程
    pattern_litigation = re.compile(r'审理经过(.*?(审理终结。|审查终结。))', re.S)
    match = pattern_litigation.search(full_text_for_litigation)
    litigation_process = match.group(1).strip() if match else "诉讼记录缺失"
    litigation_process = litigation_process.replace(fuzzy_title, "").strip()
    litigation_process = litigation_process.replace('<PARA>', '\n')
    litigation_process = "\n".join(p.strip() for p in litigation_process.split('\n') if p.strip())

    # 案件基本情况
    pos_case_info_start = full_text_for_litigation.find(litigation_process.replace('\n', ' ')) + len(litigation_process.replace('\n', ' '))
    pos_case_info_end = full_text_for_litigation.find("本院认为", pos_case_info_start)
    if pos_case_info_end == -1:
        case_info = "案件基本情况缺失"
    else:
        case_info_text = full_text_for_litigation[pos_case_info_start:pos_case_info_end].strip()
        for kw in ["一审法院认为与裁判", "二审法院认为与裁判", "再审诉讼请求", "再审辩方观点", "本院查明"]:
            case_info_text = case_info_text.replace(kw, "")
        case_info = case_info_text.replace(fuzzy_title, "").strip()
    case_info = case_info.replace('<PARA>', '\n')
    case_info = "\n".join(p.strip() for p in case_info.split('\n') if p.strip())
    print("DEBUG case_info:", repr(case_info))

    # 裁判分析过程
    pos_analysis_start = pos_case_info_end
    pos_analysis_end1 = full_text_for_litigation.find("判决如下：", pos_analysis_start)
    pos_analysis_end2 = full_text_for_litigation.find("裁定如下：", pos_analysis_start)
    if pos_analysis_end1 == -1 and pos_analysis_end2 == -1:
        trial_analysis = "裁判分析过程缺失"
    else:
        if pos_analysis_end1 == -1:
            pos_analysis_end = pos_analysis_end2
            end_marker = "裁定如下："
        elif pos_analysis_end2 == -1:
            pos_analysis_end = pos_analysis_end1
            end_marker = "判决如下："
        else:
            if pos_analysis_end1 < pos_analysis_end2:
                pos_analysis_end = pos_analysis_end1
                end_marker = "判决如下："
            else:
                pos_analysis_end = pos_analysis_end2
                end_marker = "裁定如下："
        pos_analysis_end += len(end_marker)

        trial_analysis_text = full_text_for_litigation[pos_analysis_start:pos_analysis_end].strip()
        trial_analysis = re.sub(r"本院认为", "", trial_analysis_text, count=1).strip()

    trial_analysis = trial_analysis.replace('<PARA>', '\n')
    trial_analysis = "\n".join(p.strip() for p in trial_analysis.split('\n') if p.strip())

    # 裁判结果
    pos_result_start = full_text_for_litigation.find("裁判结果", pos_analysis_end)
    pos_result_end = full_text_for_litigation.find("审判人员", pos_result_start)
    if pos_result_start == -1 or pos_result_end == -1:
        trial_result = "裁判结果缺失"
    else:
        trial_result_raw = full_text_for_litigation[pos_result_start:pos_result_end].strip()
        trial_result_raw = special_segment_trial_result(trial_result_raw)
        trial_result = trial_result_raw.replace(fuzzy_title, "").strip()
        trial_result = trial_result.replace('<PARA>', '\n')
        trial_result = "\n".join(p.strip() for p in trial_result.split('\n') if p.strip())
    if trial_result.startswith("裁判结果"):
        trial_result = trial_result[len("裁判结果"):].strip()

    # 提取审判人员信息
    judge_info = extract_judge_info_from_no_spaces("\n".join(txt_no_spaces))
    judge_info = judge_info.replace('<PARA>', '\n')
    judge_info = "\n".join(p.strip() for p in judge_info.split('\n') if p.strip())

    # 清理judge_info中汉字间多余空格（保留换行）
    judge_info_lines = judge_info.split('\n')
    cleaned_judge_lines = [re.sub(r'([\u4e00-\u9fa5])\s+([\u4e00-\u9fa5])', r'\1\2', line) for line in judge_info_lines]
    judge_info = '\n'.join(cleaned_judge_lines).strip()

    return {
        'case_name': case_name,
        'case_number': case_number,
        'litigation_process': litigation_process,
        'parties_info': parties_info,
        'case_info': case_info,
        'trial_analysis': trial_analysis,
        'trial_result': trial_result,
        'judge_info': judge_info
    }

# ---------------- 样式设置函数 ----------------

def styled_paragraph(text, color, size, bold=False, align='left', line_height=2, indent_px=16,
                     margin_top=0, margin_bottom=0, background="#ffffff",
                     font_family="'Helvetica Neue', Helvetica, 'Hiragino Sans GB', 'Microsoft YaHei', Arial, sans-serif"):
    style = f"""
        color:{color};
        font-size:{size}px;
        text-align:{align};
        line-height:{line_height};
        margin-top:{margin_top}px;
        margin-bottom:{margin_bottom}px;
        margin-left:{indent_px}px;
        margin-right:{indent_px}px;
        background-color:{background};
        font-family:{font_family};
        {'font-weight:bold;' if bold else ''}
    """
    return f'<p style="{style.strip()}">{text}</p>'


def styled_paragraphs(paragraphs, color, size, bold=False, align='left', line_height=2, indent_px=16,
                      margin_top=0, margin_bottom=0, background="#ffffff", only_last_has_margin=False,
                      last_margin_bottom=32, font_family="'Helvetica Neue', Helvetica, 'Hiragino Sans GB', 'Microsoft YaHei', Arial, sans-serif"):
    html_paras = []
    total = len(paragraphs)
    for idx, p in enumerate(paragraphs):
        is_last = (idx == total - 1)
        mb = last_margin_bottom if (only_last_has_margin and is_last) else margin_bottom
        html_paras.append(styled_paragraph(
            p, color, size, bold, align, line_height, indent_px, margin_top, mb, background, font_family
        ))
    return "\n".join(html_paras)



# ---------------- PyQt5 主程序 ----------------

class DropWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("自动排版工具")
        self.setAcceptDrops(True)
        self.resize(600, 400)
        layout = QVBoxLayout(self)

        self.label = QLabel("拖入裁判文书DOCX或PDF文件（可多选）", self)
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label)

        self.listWidget = QListWidget(self)
        layout.addWidget(self.listWidget)

        self.btn_select = QPushButton("手动选择文件", self)
        layout.addWidget(self.btn_select)
        self.btn_select.clicked.connect(self.open_file_dialog)

        self.output_dir = os.path.join(base_path, "output")
        os.makedirs(self.output_dir, exist_ok=True)

        self.listWidget.itemDoubleClicked.connect(self.open_file)

        layout.addSpacerItem(QSpacerItem(20, 20, QSizePolicy.Minimum, QSizePolicy.Expanding))
        self.author_label = QLabel("By LeClaire", self)
        self.author_label.setAlignment(Qt.AlignRight | Qt.AlignBottom)
        layout.addWidget(self.author_label)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        self.process_files(files)

    def open_file_dialog(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件", "", "文档 (*.docx *.pdf)")
        if files:
            self.process_files(files)
    
    def open_file(self, item):
        # 提取路径（兼容“✅ 处理成功：xxx（已为base64图片）”和“❌ 处理失败：xxx”）
        text = item.text()
        # 只处理“处理成功”类型
        if "处理成功" in text:
            # 提取第一个中文冒号后的路径，去掉括号内容
            try:
                path = text.split("：", 1)[1].split("（")[0].strip()
            except Exception:
                path = text  # 失败时保留原值

            if os.path.exists(path):
                if sys.platform.startswith('win'):
                    os.startfile(path)
                elif sys.platform.startswith('darwin'):
                    subprocess.call(['open', path])
                else:
                    subprocess.call(['xdg-open', path])
            else:
                self.listWidget.addItem(f"❌ 文件不存在：{path}")
        # 提取路径（兼容“✅ 处理成功：xxx（已生成base64版本）”和“❌ 处理失败：xxx”）

    def process_files(self, files):
        for path in files:
            ext = os.path.splitext(path)[1].lower()
            base_name = os.path.splitext(os.path.basename(path))[0]
            try:
                if ext == ".docx":
                    case_name = get_case_name_from_docx(path)
                    case_number = get_case_number_from_docx(path)
                    full_text = "\n".join(extract_text_from_docx(path))
                    data = parse_fields(full_text, case_name, case_number)
                elif ext == ".pdf":
                    data = extract_text_from_pdf(path, self.output_dir, base_name)
                else:
                    self.listWidget.addItem(f"❌ 非支持文件格式，跳过：{os.path.basename(path)}")
                    continue

                print("DEBUG judge_info:", repr(data['judge_info']))
                print("DEBUG parties_info:", repr(data['parties_info']))
                html_filename = f"{base_name}-公众号格式.html"
                html = generate_wechat_html(data)
                html_path = os.path.join(self.output_dir, html_filename)
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html)
                self.listWidget.addItem(f"✅ 处理成功：{html_path}（已为base64图片）")
            except Exception as e:
                err = traceback.format_exc()
                self.listWidget.addItem(f"❌ 处理失败：{os.path.basename(path)} （{e}）")
                print(f"[错误详情]\n{err}")

    # 3. generate_wechat_html 传入base64 map
def generate_wechat_html(data, image_map=None):
        parts = []
        parts.append('<meta charset="UTF-8">')
        parts.append(styled_paragraph('<span style="background-color:#5287b7;color:#ffffff;">今日案例播客版  干货知识轻松听</span>', "#5e5e5e", 16))
        parts.append(styled_paragraph('<span style="background-color:#5287b7;color:#ffffff;">延伸阅读</span>', "#5e5e5e", 16))

        parts.append(styled_paragraph('【裁判要旨】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))
        parts.append(styled_paragraph("在此输入裁判要旨内容", "#5e5e5e", 16, align='left', margin_top=0, margin_bottom=32))
        parts.append(styled_paragraph('【文书全文】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))
        parts.append(styled_paragraph('【文书标题、案号及来源】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))
    
        parts.append(styled_paragraph("标题：" + data['case_name'], "#5e5e5e", 16, margin_bottom=0))
        parts.append(styled_paragraph("案号：" + data['case_number'], "#5e5e5e", 16, margin_bottom=0))
        parts.append(styled_paragraph("来源：中国裁判文书网", "#5e5e5e", 16, margin_bottom=32))

        parts.append(styled_paragraph('【当事人信息】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))  

        parts.append(styled_paragraphs(data['parties_info'].split('\n'), "#5e5e5e", 16, margin_bottom=0, only_last_has_margin=True, last_margin_bottom=32))
        parts.append(styled_paragraph('【诉讼记录】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))  

        parts.append(styled_paragraphs(data['litigation_process'].split('\n'), "#5e5e5e", 16, margin_bottom=32))
        parts.append(styled_paragraph('【案件基本情况】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))  

        parts.append(styled_paragraphs(data['case_info'].split('\n'), "#5e5e5e", 16, margin_bottom=32, only_last_has_margin=True, last_margin_bottom=32))
        parts.append(styled_paragraph('【裁判分析过程】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))

        parts.append(styled_paragraphs(data['trial_analysis'].split('\n'), "#5e5e5e", 16, margin_bottom=32, only_last_has_margin=True, last_margin_bottom=32))
        parts.append(styled_paragraph('【裁判结果】', "#5287b7", 16, bold=True, align='left', margin_bottom=32))

        parts.append(styled_paragraphs(data['trial_result'].split('\n'), "#5e5e5e", 16, margin_bottom=0))

        parts.append('<br>' + styled_paragraphs(data['judge_info'].split('\n'), "#5e5e5e", 16, align='right', margin_bottom=0) + '<br>')

        parts.append(styled_paragraph("编辑团队：薛政  黄琳娜  初相钰  赵绮", "#5e5e5e", 16, align='center', margin_top=32, margin_bottom=0))

        return "\n".join(parts)

def convert_html_images_to_base64(html_path, output_path=None):
    """将 HTML 中本地图片路径转换为 base64"""
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    html_dir = os.path.dirname(os.path.abspath(html_path))


    output_path = output_path or html_path.replace('.html', '_base64.html')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(str(soup))

    print(f"转换完成！输出文件: {output_path}")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = DropWidget()
    win.show()
    sys.exit(app.exec_())
